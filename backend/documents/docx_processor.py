"""
StudyGemma - DOCX Processor

Responsibilities:
- Validate DOCX files
- Extract document metadata
- Extract paragraphs and text
- Detect headings
- Detect lists
- Extract tables
- Detect embedded images
- Preserve document structure
- Clean extracted content
- Perform extraction quality validation
- Prepare RAG-ready output
- Optionally send embedded images to OCR

Does NOT:
- Store files
- Generate embeddings
- Perform vector search
- Call Gemma
"""

from __future__ import annotations

import logging
import re
import zipfile
from pathlib import Path
from typing import Any, Optional, Union

try:
    from docx import Document
    from docx.document import Document as DocumentObject
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None
    DocumentObject = None
    Table = None
    _Cell = None
    Paragraph = None


logger = logging.getLogger(__name__)


class DOCXProcessor:
    """
    Processor for Microsoft Word DOCX documents.
    """

    DEFAULT_MIN_TEXT_LENGTH = 30

    def __init__(
        self,
        min_text_length: int = DEFAULT_MIN_TEXT_LENGTH,
        ocr_service: Optional[Any] = None,
    ):
        self.min_text_length = max(
            1,
            min_text_length,
        )

        self.ocr_service = ocr_service

        if Document is None:
            logger.warning(
                "python-docx is not installed. "
                "Install it with: pip install python-docx"
            )

    # =========================================================
    # MAIN PROCESS
    # =========================================================

    def process(
        self,
        docx_path: Union[str, Path],
        use_ocr: bool = True,
    ) -> dict[str, Any]:
        """
        Process a complete DOCX document.

        Returns:
            Structured document data suitable for
            further RAG processing.
        """

        try:
            path = Path(
                docx_path
            ).resolve()

            # -------------------------------------------------
            # Basic validation
            # -------------------------------------------------

            validation = self.validate_file(
                path
            )

            if not validation["valid"]:
                return self._failure(
                    validation["error"]
                )

            if Document is None:
                return self._failure(
                    "python-docx is not installed."
                )

            # -------------------------------------------------
            # Open document
            # -------------------------------------------------

            document = Document(
                str(path)
            )

            # -------------------------------------------------
            # Metadata
            # -------------------------------------------------

            metadata = self.extract_metadata(
                document,
                path,
            )

            # -------------------------------------------------
            # Extract structural content
            # -------------------------------------------------

            paragraphs = (
                self.extract_paragraphs(
                    document
                )
            )

            headings = (
                self.extract_headings(
                    paragraphs
                )
            )

            lists = (
                self.extract_lists(
                    paragraphs
                )
            )

            tables = (
                self.extract_tables(
                    document
                )
            )

            images = (
                self.extract_image_info(
                    path
                )
            )

            # -------------------------------------------------
            # OCR
            # -------------------------------------------------

            ocr_results = []

            if (
                use_ocr
                and self.ocr_service is not None
                and images
            ):

                ocr_results = (
                    self.process_images_with_ocr(
                        path
                    )
                )

            # -------------------------------------------------
            # Build clean text
            # -------------------------------------------------

            full_text = (
                self.build_full_text(
                    paragraphs=paragraphs,
                    tables=tables,
                    ocr_results=ocr_results,
                )
            )

            # -------------------------------------------------
            # Detect structure
            # -------------------------------------------------

            structure = (
                self.detect_structure(
                    paragraphs=paragraphs,
                    headings=headings,
                    lists=lists,
                    tables=tables,
                    images=images,
                )
            )

            # -------------------------------------------------
            # Quality validation
            # -------------------------------------------------

            quality = (
                self.validate_quality(
                    full_text=full_text,
                    paragraphs=paragraphs,
                    tables=tables,
                    images=images,
                )
            )

            # -------------------------------------------------
            # RAG-ready content
            # -------------------------------------------------

            rag_ready = (
                self.create_rag_ready_data(
                    metadata=metadata,
                    paragraphs=paragraphs,
                    headings=headings,
                    lists=lists,
                    tables=tables,
                    images=images,
                    ocr_results=ocr_results,
                    full_text=full_text,
                )
            )

            return {
                "success": True,
                "document": metadata,
                "paragraphs": paragraphs,
                "headings": headings,
                "lists": lists,
                "tables": tables,
                "images": images,
                "ocr_results": ocr_results,
                "full_text": full_text,
                "structure": structure,
                "quality": quality,
                "rag_ready": rag_ready,
            }

        except Exception as error:

            logger.exception(
                "DOCX processing failed."
            )

            return self._failure(
                str(error)
            )

    # =========================================================
    # VALIDATION
    # =========================================================

    @staticmethod
    def validate_file(
        path: Path,
    ) -> dict[str, Any]:
        """
        Validate that the file exists and is a valid
        DOCX/ZIP container.
        """

        if not path.exists():
            return {
                "valid": False,
                "error": (
                    f"DOCX file not found: {path}"
                ),
            }

        if not path.is_file():
            return {
                "valid": False,
                "error": (
                    "Provided path is not a file."
                ),
            }

        if path.suffix.lower() != ".docx":
            return {
                "valid": False,
                "error": (
                    "Provided file is not a DOCX file."
                ),
            }

        # -----------------------------------------------------
        # DOCX files are ZIP containers.
        # -----------------------------------------------------

        try:

            if not zipfile.is_zipfile(
                path
            ):
                return {
                    "valid": False,
                    "error": (
                        "File is not a valid DOCX "
                        "container."
                    ),
                }

            with zipfile.ZipFile(
                path,
                "r",
            ) as archive:

                names = archive.namelist()

                if (
                    "[Content_Types].xml"
                    not in names
                ):

                    return {
                        "valid": False,
                        "error": (
                            "Invalid DOCX structure."
                        ),
                    }

                if (
                    "word/document.xml"
                    not in names
                ):

                    return {
                        "valid": False,
                        "error": (
                            "Main DOCX document "
                            "content is missing."
                        ),
                    }

        except zipfile.BadZipFile:

            return {
                "valid": False,
                "error": (
                    "DOCX file is corrupted."
                ),
            }

        return {
            "valid": True,
            "error": None,
        }

    # =========================================================
    # METADATA
    # =========================================================

    @staticmethod
    def extract_metadata(
        document: Any,
        path: Path,
    ) -> dict[str, Any]:
        """
        Extract core DOCX metadata.
        """

        properties = (
            document.core_properties
        )

        return {
            "document_id": (
                DOCXProcessor.generate_document_id(
                    path
                )
            ),
            "filename": path.name,
            "path": str(path),
            "title": (
                properties.title
                or ""
            ),
            "subject": (
                properties.subject
                or ""
            ),
            "author": (
                properties.author
                or ""
            ),
            "keywords": (
                properties.keywords
                or ""
            ),
            "comments": (
                properties.comments
                or ""
            ),
            "category": (
                properties.category
                or ""
            ),
            "last_modified_by": (
                properties.last_modified_by
                or ""
            ),
            "created": (
                str(properties.created)
                if properties.created
                else None
            ),
            "modified": (
                str(properties.modified)
                if properties.modified
                else None
            ),
        }

    # =========================================================
    # PARAGRAPH EXTRACTION
    # =========================================================

    def extract_paragraphs(
        self,
        document: Any,
    ) -> list[dict[str, Any]]:
        """
        Extract paragraphs while preserving their order.
        """

        results = []

        for index, paragraph in enumerate(
            document.paragraphs
        ):

            text = self.clean_text(
                paragraph.text
            )

            if not text:
                continue

            style_name = ""

            try:
                style_name = (
                    paragraph.style.name
                    if paragraph.style
                    else ""
                )
            except Exception:
                pass

            results.append(
                {
                    "index": index,
                    "text": text,
                    "style": style_name,
                    "is_heading": (
                        self.is_heading(
                            paragraph
                        )
                    ),
                    "is_list": (
                        self.is_list(
                            paragraph
                        )
                    ),
                }
            )

        return results

    # =========================================================
    # HEADING DETECTION
    # =========================================================

    def extract_headings(
        self,
        paragraphs: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """
        Extract headings from paragraph data.
        """

        headings = []

        for paragraph in paragraphs:

            if paragraph.get(
                "is_heading",
                False,
            ):

                headings.append(
                    {
                        "index": paragraph[
                            "index"
                        ],
                        "text": paragraph[
                            "text"
                        ],
                        "style": paragraph[
                            "style"
                        ],
                    }
                )

        return headings

    # =========================================================
    # HEADING CHECK
    # =========================================================

    @staticmethod
    def is_heading(
        paragraph: Any,
    ) -> bool:
        """
        Detect common Word heading styles.
        """

        try:

            style = (
                paragraph.style.name
                or ""
            ).lower()

            if (
                "heading" in style
                or style.startswith("title")
                or style.startswith("subtitle")
            ):

                return True

        except Exception:
            pass

        # -----------------------------------------------------
        # Fallback heuristic.
        # -----------------------------------------------------

        text = (
            paragraph.text
            or ""
        ).strip()

        if not text:
            return False

        if len(text) > 120:
            return False

        if re.match(
            r"^(chapter|unit|section|module)\s+",
            text,
            re.IGNORECASE,
        ):
            return True

        return False

    # =========================================================
    # LIST DETECTION
    # =========================================================

    def extract_lists(
        self,
        paragraphs: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """
        Extract list items.
        """

        lists = []

        for paragraph in paragraphs:

            if paragraph.get(
                "is_list",
                False,
            ):

                lists.append(
                    {
                        "index": paragraph[
                            "index"
                        ],
                        "text": paragraph[
                            "text"
                        ],
                        "style": paragraph[
                            "style"
                        ],
                    }
                )

        return lists

    # =========================================================
    # LIST CHECK
    # =========================================================

    @staticmethod
    def is_list(
        paragraph: Any,
    ) -> bool:
        """
        Detect common list styles and textual lists.
        """

        try:

            style = (
                paragraph.style.name
                or ""
            ).lower()

            if (
                "list" in style
                or "bullet" in style
                or "number" in style
            ):

                return True

        except Exception:
            pass

        text = (
            paragraph.text
            or ""
        ).strip()

        if re.match(
            r"^[-*•]\s+",
            text,
        ):
            return True

        if re.match(
            r"^\d+[\.\)]\s+",
            text,
        ):
            return True

        return False

    # =========================================================
    # TABLE EXTRACTION
    # =========================================================

    def extract_tables(
        self,
        document: Any,
    ) -> list[dict[str, Any]]:
        """
        Extract DOCX tables into structured rows/cells.
        """

        results = []

        for table_index, table in enumerate(
            document.tables
        ):

            rows = []

            for row_index, row in enumerate(
                table.rows
            ):

                cells = []

                for column_index, cell in enumerate(
                    row.cells
                ):

                    cell_text = self.clean_text(
                        cell.text
                    )

                    cells.append(
                        {
                            "column": column_index,
                            "text": cell_text,
                        }
                    )

                rows.append(
                    {
                        "row": row_index,
                        "cells": cells,
                    }
                )

            results.append(
                {
                    "table_index": table_index,
                    "rows": rows,
                    "row_count": len(
                        rows
                    ),
                    "column_count": (
                        max(
                            (
                                len(row["cells"])
                                for row in rows
                            ),
                            default=0,
                        )
                    ),
                }
            )

        return results

    # =========================================================
    # IMAGE INFORMATION
    # =========================================================

    @staticmethod
    def extract_image_info(
        docx_path: Path,
    ) -> list[dict[str, Any]]:
        """
        Detect embedded images in the DOCX package.

        Images are not permanently extracted here.
        """

        images = []

        try:

            with zipfile.ZipFile(
                docx_path,
                "r",
            ) as archive:

                media_files = [
                    name
                    for name in archive.namelist()
                    if name.startswith(
                        "word/media/"
                    )
                ]

                for index, name in enumerate(
                    media_files
                ):

                    info = archive.getinfo(
                        name
                    )

                    extension = (
                        Path(name)
                        .suffix
                        .lower()
                    )

                    images.append(
                        {
                            "index": index + 1,
                            "name": Path(
                                name
                            ).name,
                            "path_in_document": name,
                            "extension": extension,
                            "size": info.file_size,
                        }
                    )

        except Exception as error:

            logger.warning(
                "Could not inspect DOCX images: %s",
                error,
            )

        return images

    # =========================================================
    # OCR IMAGES
    # =========================================================

    def process_images_with_ocr(
        self,
        docx_path: Path,
    ) -> list[dict[str, Any]]:
        """
        Extract embedded images temporarily in memory
        and send them to OCR.

        The images are not permanently written to disk.
        """

        if self.ocr_service is None:

            return []

        results = []

        try:

            with zipfile.ZipFile(
                docx_path,
                "r",
            ) as archive:

                media_files = [
                    name
                    for name in archive.namelist()
                    if name.startswith(
                        "word/media/"
                    )
                ]

                for index, name in enumerate(
                    media_files
                ):

                    try:

                        image_bytes = (
                            archive.read(
                                name
                            )
                        )

                        text = (
                            self.run_ocr(
                                image_bytes
                            )
                        )

                        if text:

                            results.append(
                                {
                                    "image_index": (
                                        index + 1
                                    ),
                                    "image_name": (
                                        Path(
                                            name
                                        ).name
                                    ),
                                    "text": (
                                        self.clean_text(
                                            text
                                        )
                                    ),
                                }
                            )

                    except Exception as error:

                        logger.warning(
                            "OCR failed for image %s: %s",
                            name,
                            error,
                        )

        except Exception as error:

            logger.warning(
                "Could not process DOCX images: %s",
                error,
            )

        return results

    # =========================================================
    # OCR CALL
    # =========================================================

    def run_ocr(
        self,
        image_bytes: bytes,
    ) -> str:
        """
        Call the configured OCR service.
        """

        if self.ocr_service is None:

            return ""

        try:

            if hasattr(
                self.ocr_service,
                "extract_text",
            ):

                result = (
                    self.ocr_service.extract_text(
                        image_bytes
                    )
                )

            elif hasattr(
                self.ocr_service,
                "process_image",
            ):

                result = (
                    self.ocr_service.process_image(
                        image_bytes
                    )
                )

            elif callable(
                self.ocr_service
            ):

                result = (
                    self.ocr_service(
                        image_bytes
                    )
                )

            else:

                return ""

            if isinstance(
                result,
                str,
            ):

                return result

            if isinstance(
                result,
                dict,
            ):

                return str(
                    result.get(
                        "text",
                        "",
                    )
                )

            return str(
                result or ""
            )

        except Exception as error:

            logger.warning(
                "OCR service failed: %s",
                error,
            )

            return ""

    # =========================================================
    # BUILD FULL TEXT
    # =========================================================

    def build_full_text(
        self,
        paragraphs: list[dict[str, Any]],
        tables: list[dict[str, Any]],
        ocr_results: list[dict[str, Any]],
    ) -> str:
        """
        Build a unified text representation while
        preserving basic document structure.
        """

        sections = []

        # -----------------------------------------------------
        # Paragraphs
        # -----------------------------------------------------

        for paragraph in paragraphs:

            text = paragraph.get(
                "text",
                "",
            )

            if not text:
                continue

            if paragraph.get(
                "is_heading",
                False,
            ):

                sections.append(
                    f"\n## {text}\n"
                )

            elif paragraph.get(
                "is_list",
                False,
            ):

                sections.append(
                    f"- {text}"
                )

            else:

                sections.append(
                    text
                )

        # -----------------------------------------------------
        # Tables
        # -----------------------------------------------------

        for table in tables:

            sections.append(
                "\n[TABLE]"
            )

            for row in table.get(
                "rows",
                [],
            ):

                cell_values = [
                    cell.get(
                        "text",
                        "",
                    )
                    for cell in row.get(
                        "cells",
                        [],
                    )
                ]

                sections.append(
                    " | ".join(
                        cell_values
                    )
                )

            sections.append(
                "[/TABLE]"
            )

        # -----------------------------------------------------
        # OCR text
        # -----------------------------------------------------

        for result in ocr_results:

            text = result.get(
                "text",
                "",
            )

            if text:

                sections.append(
                    f"\n[IMAGE TEXT]\n{text}"
                )

        return self.clean_text(
            "\n".join(
                sections
            )
        )

    # =========================================================
    # STRUCTURE DETECTION
    # =========================================================

    @staticmethod
    def detect_structure(
        paragraphs: list[dict[str, Any]],
        headings: list[dict[str, Any]],
        lists: list[dict[str, Any]],
        tables: list[dict[str, Any]],
        images: list[dict[str, Any]],
    ) -> dict[str, Any]:
        """
        Produce a high-level document structure summary.
        """

        return {
            "paragraph_count": len(
                paragraphs
            ),
            "heading_count": len(
                headings
            ),
            "list_item_count": len(
                lists
            ),
            "table_count": len(
                tables
            ),
            "image_count": len(
                images
            ),
            "has_headings": bool(
                headings
            ),
            "has_lists": bool(
                lists
            ),
            "has_tables": bool(
                tables
            ),
            "has_images": bool(
                images
            ),
        }

    # =========================================================
    # QUALITY VALIDATION
    # =========================================================

    def validate_quality(
        self,
        full_text: str,
        paragraphs: list[dict[str, Any]],
        tables: list[dict[str, Any]],
        images: list[dict[str, Any]],
    ) -> dict[str, Any]:
        """
        Validate extracted DOCX content.
        """

        text_length = len(
            full_text.strip()
        )

        has_text = (
            text_length
            >= self.min_text_length
        )

        if not full_text.strip():

            status = "poor"

        elif not has_text:

            status = "weak"

        else:

            status = "good"

        return {
            "status": status,
            "has_text": bool(
                full_text.strip()
            ),
            "text_sufficient": has_text,
            "text_length": text_length,
            "paragraph_count": len(
                paragraphs
            ),
            "table_count": len(
                tables
            ),
            "image_count": len(
                images
            ),
        }

    # =========================================================
    # RAG-READY DATA
    # =========================================================

    def create_rag_ready_data(
        self,
        metadata: dict[str, Any],
        paragraphs: list[dict[str, Any]],
        headings: list[dict[str, Any]],
        lists: list[dict[str, Any]],
        tables: list[dict[str, Any]],
        images: list[dict[str, Any]],
        ocr_results: list[dict[str, Any]],
        full_text: str,
    ) -> dict[str, Any]:
        """
        Create a structured representation for the RAG layer.

        Chunking and embeddings happen later.
        """

        blocks = []

        # -----------------------------------------------------
        # Paragraph blocks
        # -----------------------------------------------------

        for paragraph in paragraphs:

            block_type = "paragraph"

            if paragraph.get(
                "is_heading",
                False,
            ):

                block_type = "heading"

            elif paragraph.get(
                "is_list",
                False,
            ):

                block_type = "list"

            blocks.append(
                {
                    "type": block_type,
                    "index": paragraph[
                        "index"
                    ],
                    "text": paragraph[
                        "text"
                    ],
                    "metadata": {
                        "style": paragraph.get(
                            "style",
                            "",
                        ),
                    },
                }
            )

        # -----------------------------------------------------
        # Table blocks
        # -----------------------------------------------------

        for table in tables:

            table_text = []

            for row in table.get(
                "rows",
                [],
            ):

                values = [
                    cell.get(
                        "text",
                        "",
                    )
                    for cell in row.get(
                        "cells",
                        [],
                    )
                ]

                table_text.append(
                    " | ".join(
                        values
                    )
                )

            blocks.append(
                {
                    "type": "table",
                    "index": table[
                        "table_index"
                    ],
                    "text": "\n".join(
                        table_text
                    ),
                    "metadata": {
                        "row_count": table[
                            "row_count"
                        ],
                        "column_count": table[
                            "column_count"
                        ],
                    },
                }
            )

        # -----------------------------------------------------
        # OCR blocks
        # -----------------------------------------------------

        for result in ocr_results:

            blocks.append(
                {
                    "type": "image_text",
                    "index": result.get(
                        "image_index"
                    ),
                    "text": result.get(
                        "text",
                        "",
                    ),
                    "metadata": {
                        "image_name": result.get(
                            "image_name",
                            "",
                        ),
                        "ocr": True,
                    },
                }
            )

        return {
            "document_id": metadata[
                "document_id"
            ],
            "source_type": "docx",
            "filename": metadata[
                "filename"
            ],
            "text": full_text,
            "blocks": blocks,
            "metadata": {
                "title": metadata.get(
                    "title",
                    "",
                ),
                "author": metadata.get(
                    "author",
                    "",
                ),
                "headings": headings,
                "image_count": len(
                    images
                ),
                "table_count": len(
                    tables
                ),
            },
        }

    # =========================================================
    # CLEAN TEXT
    # =========================================================

    @staticmethod
    def clean_text(
        text: Optional[str],
    ) -> str:
        """
        Clean extracted DOCX text without destroying
        meaningful content.
        """

        if not text:

            return ""

        text = text.replace(
            "\r\n",
            "\n",
        )

        text = text.replace(
            "\r",
            "\n",
        )

        text = text.replace(
            "\x00",
            "",
        )

        # Normalize spaces but preserve newlines.
        text = re.sub(
            r"[ \t]+",
            " ",
            text,
        )

        # Remove spaces around line breaks.
        text = re.sub(
            r" *\n *",
            "\n",
            text,
        )

        # Avoid excessive blank lines.
        text = re.sub(
            r"\n{3,}",
            "\n\n",
            text,
        )

        return text.strip()

    # =========================================================
    # DOCUMENT ID
    # =========================================================

    @staticmethod
    def generate_document_id(
        path: Path,
    ) -> str:
        """
        Generate a deterministic document identifier
        from the resolved file path.
        """

        import hashlib

        value = str(
            path
        ).encode(
            "utf-8"
        )

        return hashlib.sha256(
            value
        ).hexdigest()[:32]

    # =========================================================
    # SIMPLE TEXT EXTRACTION
    # =========================================================

    def extract_text(
        self,
        docx_path: Union[str, Path],
    ) -> str:
        """
        Convenience method when only text is needed.
        """

        result = self.process(
            docx_path,
            use_ocr=False,
        )

        if not result.get(
            "success",
            False,
        ):

            raise RuntimeError(
                result.get(
                    "error",
                    "DOCX processing failed.",
                )
            )

        return result.get(
            "full_text",
            "",
        )

    # =========================================================
    # FAILURE RESPONSE
    # =========================================================

    @staticmethod
    def _failure(
        error: str,
    ) -> dict[str, Any]:
        """
        Standard error response.
        """

        return {
            "success": False,
            "error": error,
        }


# =============================================================
# DEFAULT INSTANCE
# =============================================================

docx_processor = DOCXProcessor()
